"""
Extract challenge documents (docx -> txt) for analysis.
"""
import os
import sys
from docx import Document

DATA_DIR = "data"
OUTPUT_DIR = "data"

docs = [
    "README.docx",
    "job_description.docx",
    "submission_spec.docx",
    "redrob_signals_doc.docx",
]

for fname in docs:
    path = os.path.join(DATA_DIR, fname)
    outpath = os.path.join(OUTPUT_DIR, fname.replace(".docx", ".txt"))

    if not os.path.exists(path):
        print(f"SKIP: {fname} not found")
        continue

    doc = Document(path)
    text = "\n".join([p.text for p in doc.paragraphs])

    with open(outpath, "w", encoding="utf-8") as f:
        f.write(text)

    print(f"EXTRACTED: {fname} -> {outpath} ({len(text)} chars, {len(doc.paragraphs)} paragraphs)")

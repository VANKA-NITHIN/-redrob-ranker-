"""
Deep extraction of docx files - including tables and all content.
"""
import os
from docx import Document

path = os.path.join("data", "redrob_signals_doc.docx")
doc = Document(path)

print("=" * 60)
print("PARAGRAPHS:")
print("=" * 60)
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[P{i}] {p.text}")

print("\n" + "=" * 60)
print("TABLES:")
print("=" * 60)
for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  Row {ri}: {' | '.join(cells)}")

print("\n" + "=" * 60)
print("ALL TEXT (including tables inline):")
print("=" * 60)
# Combine paragraphs and tables
all_text = []
for p in doc.paragraphs:
    if p.text.strip():
        all_text.append(p.text)
for ti, table in enumerate(doc.tables):
    all_text.append(f"\n[TABLE {ti}]")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        all_text.append(" | ".join(cells))

full = "\n".join(all_text)
print(f"Total extracted: {len(full)} chars")
print(full)

# Also check for other content like images, headers, footers
print("\n" + "=" * 60)
print("SECTIONS:")
print("=" * 60)
for si, section in enumerate(doc.sections):
    print(f"Section {si}:")
    header = section.header
    footer = section.footer
    if header.paragraphs:
        for p in header.paragraphs:
            if p.text.strip():
                print(f"  Header: {p.text}")
    if footer.paragraphs:
        for p in footer.paragraphs:
            if p.text.strip():
                print(f"  Footer: {p.text}")
